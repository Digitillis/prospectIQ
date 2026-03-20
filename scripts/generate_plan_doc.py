"""Generate a Word document from the ProspectIQ design plan."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def main():
    doc = Document()

    # Title
    title = doc.add_heading('ProspectIQ System Design Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Digitillis Sales Intelligence & Engagement Platform', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Version 1.0 | February 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # Context
    add_heading(doc, 'Context', 1)
    doc.add_paragraph(
        'Digitillis is a pre-revenue AI-native manufacturing intelligence platform, initially seeking '
        '3 pilot customers (6-8 week pilots), and then building a pipeline of credible prospects and '
        'then paying customers in Midwest US discrete manufacturing. ProspectIQ is a standalone AI-powered '
        'sales intelligence and engagement system to systematically identify, research, qualify, and engage '
        'manufacturing prospects with deeply personalized, multi-stage outreach.'
    )
    doc.add_paragraph('Goal: First qualified outreach within ~2 weeks of starting development.').bold = True

    # Target Market
    add_heading(doc, 'Target Market', 2)
    add_table(doc, ['Tier', 'Sub-Sector', 'NAICS Prefix'], [
        ['1A', 'Industrial Machinery & Heavy Equipment', '333'],
        ['1B', 'Automotive Parts & Components', '336'],
        ['2', 'Metal Fabrication & Precision Machining', '332'],
        ['3', 'Plastics & Injection Molding', '326'],
        ['4', 'Electronics Assembly & Semiconductor', '334'],
        ['5', 'Aerospace Components', '3364'],
    ])
    doc.add_paragraph('')
    doc.add_paragraph('Geography: Midwest US (IL, IN, MI, OH, WI, MN, IA, MO)')
    doc.add_paragraph('Company Revenue: $500M - $8B')
    doc.add_paragraph('Buyer Personas: VP Operations, Plant Manager, Head of Digital Transformation, COO, CIO')

    # Tool Stack
    add_heading(doc, 'Confirmed Tool Stack', 2)
    add_table(doc, ['Component', 'Tool', 'Est. Monthly Cost'], [
        ['Lead Source', 'Apollo.io (Professional)', '$99'],
        ['CRM / Data Layer', 'Supabase (PostgreSQL)', '$25'],
        ['AI Agents', 'Python + Claude API (Anthropic)', '$30-80'],
        ['Research', 'Perplexity API (sonar-pro)', '$5-20'],
        ['Cold Email', 'Instantly.ai (warmup + sequences)', '$30'],
        ['Transactional Email', 'Resend', '$20'],
        ['Frontend', 'Next.js on Vercel', 'Free'],
        ['Backend Hosting', 'Railway', '$0-20'],
        ['Total', '', '$210-275/mo'],
    ])

    # Project Structure
    add_heading(doc, '1. Project Structure', 1)
    structure = """prospectIQ/
├── backend/
│   ├── app/
│   │   ├── agents/          # 7 AI agent modules
│   │   ├── integrations/    # Apollo, Perplexity, Instantly, Resend clients
│   │   ├── api/routes/      # FastAPI endpoints (serves dashboard)
│   │   ├── core/            # Config loader, DB client, models, cost tracker
│   │   ├── orchestrator/    # Pipeline coordinator + state machine
│   │   └── utils/           # Territory mapping, NAICS utilities
│   └── scripts/             # CLI entry points (run_discovery, run_research, etc.)
├── dashboard/               # Next.js CRM frontend (Vercel)
│   └── app/                 # Pipeline, Prospects, Approvals, Actions, Analytics pages
├── config/                  # YAML configuration (version controlled)
│   ├── icp.yaml             # Ideal Customer Profile filters
│   ├── scoring.yaml         # PQS weights, signals, thresholds
│   ├── sequences.yaml       # Engagement sequence definitions
│   └── manufacturing_ontology.yaml  # NAICS, tech systems, pain points, competitors
└── supabase/migrations/     # Database schema"""
    doc.add_paragraph(structure, style='No Spacing')

    # Database Schema
    add_heading(doc, '2. Database Schema', 1)
    doc.add_paragraph('8 core tables in Supabase (PostgreSQL):')

    add_heading(doc, 'companies — Core company records', 3)
    add_table(doc, ['Column Group', 'Fields', 'Notes'], [
        ['Identity', 'id (UUID PK), apollo_id, name, domain, website', 'apollo_id for deduplication'],
        ['Industry', 'industry, naics_code, sub_sector, tier', 'Tier: 1a/1b/2/3/4/5'],
        ['Firmographic', 'employee_count, revenue_range, estimated_revenue, founded_year, is_private', ''],
        ['Location', 'city, state, country, territory', 'Territory = deterministic from state'],
        ['AI Intelligence', 'research_summary, technology_stack (JSONB), pain_signals (JSONB), manufacturing_profile (JSONB), personalization_hooks (JSONB)', 'Populated by Research Agent'],
        ['Qualification', 'pqs_total, pqs_firmographic, pqs_technographic, pqs_timing, pqs_engagement', '0-25 each dimension, 0-100 total'],
        ['Lifecycle', 'status (enum), campaign_name, batch_id', '15 lifecycle stages'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'contacts — Individual contacts at companies', 3)
    add_table(doc, ['Column Group', 'Fields'], [
        ['Identity', 'id, company_id (FK), apollo_id, first_name, last_name, email, phone'],
        ['Professional', 'title, seniority, department, headline, linkedin_url'],
        ['Classification', 'is_decision_maker (bool), persona_type (vp_ops/plant_manager/cio/coo/digital_transformation)'],
        ['Status', 'status (identified → enriched → contacted → engaged → not_interested → bounced)'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'research_intelligence — AI research per company (1:1)', 3)
    add_table(doc, ['Column Group', 'Fields'], [
        ['Raw Data', 'perplexity_response, claude_analysis'],
        ['Manufacturing Profile', 'manufacturing_type, equipment_types (JSONB), known_systems (JSONB)'],
        ['Maturity', 'iot_maturity (none/basic/intermediate/advanced), maintenance_approach (reactive/time_based/condition_based/predictive)'],
        ['Intelligence', 'pain_points (JSONB), opportunities (JSONB), existing_solutions (JSONB)'],
        ['Financial', 'funding_status, funding_details'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'Other Tables', 3)
    add_table(doc, ['Table', 'Purpose', 'Key Fields'], [
        ['outreach_drafts', 'AI-generated messages pending approval', 'subject, body, approval_status, personalization_notes'],
        ['interactions', 'Every touchpoint (immutable log)', 'type, channel, metadata (JSONB), source'],
        ['engagement_sequences', 'Active sequences per contact', 'sequence_name, current_step, next_action_at, status'],
        ['api_costs', 'Cost tracking per API call', 'provider, model, input_tokens, output_tokens, estimated_cost_usd'],
        ['learning_outcomes', 'What worked/didn\'t', 'message_theme, outcome, company_tier, sub_sector, persona_type'],
    ])

    # Agent Specifications
    add_heading(doc, '3. Agent Specifications', 1)
    doc.add_paragraph('7 specialized Python agents, each with clear input/output contracts:')

    add_table(doc, ['Agent', 'Purpose', 'LLM Calls', 'Input', 'Output'], [
        ['1. Discovery', 'Search Apollo for companies + contacts matching ICP', 'None (API only)', 'icp.yaml', 'companies + contacts in Supabase'],
        ['2. Research', 'Deep-research each company for tech, timing, pain signals', '1 Perplexity + 1 Claude per company', 'Discovered companies', 'research_intelligence records'],
        ['3. Qualification', 'Score all 4 PQS dimensions, classify leads', 'None (rule-based from config)', 'Researched companies', 'PQS scores, status changes'],
        ['4. Outreach', 'Generate personalized outreach messages', '1 Claude per message', 'Qualified companies + research', 'outreach_drafts (pending approval)'],
        ['5. Engagement', 'Orchestrate multi-stage sequences, handle events', '1 Claude per follow-up', 'Approved outreach + webhooks', 'Sent emails, follow-up drafts'],
        ['6. Reply', 'Classify replies and draft responses', '1 Claude per reply', 'Instantly.ai webhooks', 'Classification + response drafts'],
        ['7. Learning', 'Analyze outcomes, surface insights', '1 Claude weekly', 'learning_outcomes data', 'Weekly summary, suggestions'],
    ])

    doc.add_paragraph('')

    # Discovery Agent detail
    add_heading(doc, 'Agent 1: Discovery Agent', 2)
    doc.add_paragraph('Uses Apollo People Search API (free, no credits consumed) to find contacts matching ICP filters. '
                      'Deduplicates against existing Supabase records. Auto-classifies persona_type from title. '
                      'Calculates initial firmographic PQS score. No LLM calls needed.')

    # Research Agent detail
    add_heading(doc, 'Agent 2: Research Agent', 2)
    doc.add_paragraph('Two-call pipeline per company (down from 5 in reference implementation):')
    p = doc.add_paragraph()
    p.add_run('Call 1 — Perplexity (sonar-pro): ').bold = True
    p.add_run('Manufacturing-specific research prompt covering products, technology systems, IoT maturity, '
              'recent news, sustainability, maintenance approach, AI/ML platforms, hiring signals.')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Call 2 — Claude (Sonnet): ').bold = True
    p.add_run('Structured analysis extracting technology_stack, pain_signals, equipment_types, known_systems, '
              'iot_maturity, maintenance_approach, personalization_hooks (3-5 specific facts for outreach), '
              'opportunities, existing_solutions, confidence_level. Output as JSON.')
    doc.add_paragraph('')
    doc.add_paragraph('Cost per company: ~$0.04-0.08')

    # Qualification Agent detail
    add_heading(doc, 'Agent 3: Qualification Agent', 2)
    doc.add_paragraph('Rule-based scoring driven entirely by scoring.yaml configuration. No LLM calls needed.')

    # Outreach Agent detail
    add_heading(doc, 'Agent 4: Outreach Agent', 2)
    doc.add_paragraph('Prompt design principles:')
    bullets = [
        'Lead with the prospect\'s specific challenge, not Digitillis features',
        'Reference concrete facts from research (personalization hooks)',
        'Short: 100-150 words max for email body',
        'Single clear CTA per message',
        'Value-first: offer genuine insight, not just a meeting request',
        'Match tone to persona (technical for CIO, business-outcome for COO)',
        'No snake-oil, no urgency manipulation, no false scarcity',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Qualification Framework
    add_heading(doc, '4. Qualification Framework (PQS)', 1)
    doc.add_paragraph('Prospect Quality Score: 4 dimensions, 0-25 points each, total 0-100.')

    add_heading(doc, 'Dimension 1: Firmographic Fit (0-25 pts)', 2)
    add_table(doc, ['Signal', 'Points', 'Source'], [
        ['Discrete manufacturing (NAICS 31-33)', '5', 'Apollo'],
        ['Revenue $500M-$8B', '5', 'Apollo'],
        ['Midwest US', '5', 'Apollo'],
        ['Employee count 500-20,000', '3', 'Apollo'],
        ['Multiple plant locations', '4', 'Research'],
        ['Private company', '3', 'Apollo'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'Dimension 2: Technographic Readiness (0-25 pts)', 2)
    add_table(doc, ['Signal', 'Points', 'Source'], [
        ['Uses legacy CMMS/APM (SAP PM, Maximo, UpKeep, Fiix)', '5', 'Research'],
        ['Has IoT/sensor infrastructure', '5', 'Research'],
        ['ERP in place (SAP, Oracle, Epicor, Infor)', '4', 'Research'],
        ['OPC-UA / MQTT / Modbus infrastructure', '4', 'Research'],
        ['No existing AI/ML platform for manufacturing', '4', 'Research'],
        ['Industry 4.0 / smart factory initiative', '3', 'Research'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'Dimension 3: Timing & Pain Signals (0-25 pts)', 2)
    add_table(doc, ['Signal', 'Points', 'Source'], [
        ['Recently hired digital transformation role', '5', 'Research'],
        ['Unplanned downtime mentioned', '5', 'Research'],
        ['Recent capex on equipment / plant expansion', '4', 'Research'],
        ['Sustainability / ESG mandates', '3', 'Research'],
        ['Quality issues (recalls, complaints)', '3', 'Research'],
        ['Workforce challenges (skilled labor shortage)', '3', 'Research'],
        ['Recent M&A (integrating new plants)', '2', 'Research'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'Dimension 4: Engagement Progress (0-25 pts)', 2)
    add_table(doc, ['Stage', 'Points', 'Trigger'], [
        ['Cold (no interaction)', '0', 'Default'],
        ['Delivered (outreach sent)', '2', 'Email sent'],
        ['Opened / Viewed', '5', 'Email open / profile view'],
        ['Engaged (clicked, replied)', '10', 'Any response'],
        ['Interested (requested info)', '15', 'Explicit interest'],
        ['Evaluating (pilot discussion)', '20', 'Meeting scheduled'],
        ['Committed (pilot signed)', '25', 'Agreement signed'],
    ])

    doc.add_paragraph('')
    add_heading(doc, 'Qualification Thresholds', 2)
    add_table(doc, ['PQS Range', 'Classification', 'Action'], [
        ['0-25', 'Unqualified', 'Do not pursue'],
        ['26-45', 'Research Needed', 'Invest in deep research to determine fit'],
        ['46-60', 'Qualified Lead', 'Proceed with personalized outreach'],
        ['61-75', 'High-Priority Lead', 'Accelerate outreach, multi-channel approach'],
        ['76-100', 'Hot Prospect', 'Immediate attention, founder-direct outreach'],
    ])

    # Engagement Design
    add_heading(doc, '5. Engagement Sequences', 1)

    add_heading(doc, 'Initial Outreach Sequence (5 steps)', 2)
    add_table(doc, ['Step', 'Channel', 'Delay', 'Approach'], [
        ['1', 'Email', 'Day 0', 'Lead with specific pain point, introduce Digitillis (150 words max)'],
        ['2', 'LinkedIn', 'Day 2', 'Connection request with brief value reference (manual)'],
        ['3', 'Email', 'Day 5', 'Share relevant industry insight or data point (120 words)'],
        ['4', 'Email', 'Day 7', 'Reference similar company challenge, describe Digitillis solution (130 words)'],
        ['5', 'Email', 'Day 10', 'Direct but respectful ask — offer assessment/demo/ROI estimate (100 words)'],
    ])

    doc.add_paragraph('')
    doc.add_paragraph('Additional sequences: Warm Follow-up (for opens without reply), Positive Reply (next steps)')

    # Lifecycle State Machine
    add_heading(doc, 'Lifecycle State Machine', 2)
    doc.add_paragraph(
        'DISCOVERED → RESEARCHED → QUALIFIED → OUTREACH_PENDING → CONTACTED → '
        'ENGAGED → MEETING_SCHEDULED → PILOT_DISCUSSION → PILOT_SIGNED → ACTIVE_PILOT → CONVERTED'
    )
    doc.add_paragraph('Terminal states: DISQUALIFIED, NOT_INTERESTED, BOUNCED, PAUSED')

    # CRM Dashboard
    add_heading(doc, '6. CRM Dashboard Design', 1)
    doc.add_paragraph('Next.js 15 + Tailwind CSS + shadcn/ui, deployed on Vercel')

    add_table(doc, ['Page', 'Purpose', 'Key Features'], [
        ['/ (Pipeline)', 'Pipeline overview', 'Kanban board by status, key metrics, pending approval count'],
        ['/prospects', 'Prospect list', 'Sortable/filterable table, PQS scores, tier badges'],
        ['/prospects/[id]', 'Prospect detail', 'Research intelligence, PQS breakdown, contacts, interaction timeline, outreach history'],
        ['/approvals', 'Approval queue', 'Pending messages with company context, approve/edit/reject actions'],
        ['/actions', 'Daily action list', 'Follow-ups due, LinkedIn touches, hot replies, new qualifications'],
        ['/analytics', 'Performance', 'Outreach funnel, metrics by segment, API cost tracking'],
    ])

    # Human in the Loop
    add_heading(doc, '7. Human-in-the-Loop Workflow', 1)
    doc.add_paragraph('All external communications require founder approval:')
    bullets = [
        'Review and approve/reject qualified leads before outreach',
        'Review and approve/edit AI-drafted outreach messages before sending',
        'Review and approve follow-up messages',
        'Review and respond to incoming replies',
        'All state transitions involving external communication need approval',
        'Dashboard surfaces daily action queue of items needing attention',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Deployment
    add_heading(doc, '8. Deployment Architecture', 1)
    add_table(doc, ['Component', 'Platform', 'Notes'], [
        ['Next.js Dashboard', 'Vercel', 'Auto-deploy from git'],
        ['Supabase (PostgreSQL)', 'Supabase Cloud', 'Managed database + realtime'],
        ['Python Backend (FastAPI)', 'Railway (later)', 'Run locally initially, move to Railway for webhooks'],
        ['Agent Scripts', 'Local CLI', 'Run manually, then schedule on Railway'],
    ])

    doc.add_paragraph('')
    doc.add_paragraph('Initial approach: Run Python agents locally via CLI scripts, deploy dashboard to Vercel. '
                      'Move to Railway when webhook handling and scheduled runs are needed.')

    # Build Timeline
    add_heading(doc, '9. Build Timeline', 1)
    add_table(doc, ['Phase', 'Component', 'Duration'], [
        ['1', 'Supabase schema + config files', 'Day 1'],
        ['2', 'Core utilities (config loader, DB client, models)', 'Day 2'],
        ['3', 'Apollo client + Discovery Agent', 'Day 2-3'],
        ['4', 'Perplexity client + Research Agent', 'Day 3-4'],
        ['5', 'Qualification Agent', 'Day 4-5'],
        ['MILESTONE', 'First batch of qualified prospects', 'Day 5'],
        ['6', 'Outreach Agent', 'Day 6'],
        ['7', 'FastAPI backend', 'Day 6-7'],
        ['8', 'Next.js CRM dashboard', 'Day 7-10'],
        ['9', 'Instantly.ai client + Engagement Agent', 'Day 8-10'],
        ['MILESTONE', 'First outreach sent', 'Day 10-12'],
        ['10', 'Reply Agent + Daily Actions', 'Day 11-13'],
        ['11', 'Analytics + Learning Agent', 'Day 13-15'],
    ])

    # Cost Estimate
    add_heading(doc, '10. Monthly Cost Estimate', 1)
    add_table(doc, ['Service', 'Monthly Cost'], [
        ['Apollo.io Professional', '$99'],
        ['Supabase Pro', '$25'],
        ['Anthropic API (Claude)', '$30-80'],
        ['Perplexity API', '$5-20'],
        ['Instantly.ai', '$30'],
        ['Resend', '$20'],
        ['Vercel', 'Free tier'],
        ['Railway', '$0-20'],
        ['Total', '$210-275/mo'],
    ])

    # Save
    output_path = os.path.expanduser('/Users/avanish/prospectIQ/ProspectIQ_Design_Plan.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == '__main__':
    main()
