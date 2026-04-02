"""Generate ProspectIQ Product Design Document as a Word file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=None, color=None, space_before=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells, header=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if header:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if bg_color:
            set_cell_background(cell, bg_color)
    return row

def style_table(table):
    """Apply alternating row styling."""
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_background(cell, '1E293B')
            elif i % 2 == 0:
                set_cell_background(cell, 'F8FAFC')

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

# ─── Styles ──────────────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(10.5)

# ─── COVER PAGE ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('ProspectIQ')
title_run.bold = True
title_run.font.size = Pt(42)
title_run.font.color.rgb = RGBColor(15, 118, 110)  # teal-700

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run('AI-Native B2B Sales Intelligence Platform')
sub_run.font.size = Pt(18)
sub_run.font.color.rgb = RGBColor(71, 85, 105)  # slate-500
sub_run.bold = True

doc.add_paragraph()

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_run = tagline_p.add_run('Market Analysis · Gap Assessment · Product Design Document · Implementation Plan')
tag_run.font.size = Pt(11)
tag_run.italic = True
tag_run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f'Prepared: {datetime.date.today().strftime("%B %Y")}')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(148, 163, 184)

version_p = doc.add_paragraph()
version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_run = version_p.add_run('Version 1.0 — Internal Strategy Document')
ver_run.font.size = Pt(10)
ver_run.font.color.rgb = RGBColor(148, 163, 184)

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf_run = conf_p.add_run('CONFIDENTIAL')
conf_run.bold = True
conf_run.font.size = Pt(10)
conf_run.font.color.rgb = RGBColor(220, 38, 38)

doc.add_page_break()

# ─── TABLE OF CONTENTS ───────────────────────────────────────────────────────
add_heading(doc, 'Table of Contents', level=1)
toc_items = [
    ('1', 'Executive Summary', '3'),
    ('2', 'Market Opportunity Analysis', '4'),
    ('3', 'Competitive Landscape & Positioning', '6'),
    ('4', 'Pricing & Valuation Model', '8'),
    ('5', 'Gap Analysis: Current State vs Market-Ready', '10'),
    ('6', 'Product Vision & Feature Design', '13'),
    ('7', 'User Experience Design', '20'),
    ('8', 'Effort Estimation', '24'),
    ('9', 'Implementation Plan', '26'),
    ('10', 'Risk Register', '29'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num}.  {title}').bold = False
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Executive Summary', level=1)

add_para(doc,
    'ProspectIQ is an AI-native B2B sales intelligence platform built specifically for '
    'identifying, researching, qualifying, and engaging manufacturing companies at scale. '
    'What began as an internal go-to-market engine for Digitillis has the architecture, '
    'depth, and differentiation to stand as an independent commercial product.',
    space_after=8)

add_para(doc,
    'This document provides: (1) a rigorous market opportunity analysis with pricing and '
    'valuation models, (2) an honest gap assessment of what must be built to reach market '
    'readiness, (3) a detailed product design specification including UX flows and feature '
    'definitions, and (4) a phased implementation plan with effort estimates.',
    space_after=8)

add_heading(doc, 'Key Findings at a Glance', level=2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Dimension'
hdr[1].text = 'Finding'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(cell, '1E293B')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

rows_data = [
    ('Market Size', 'TAM $4.2B (B2B data + sales intelligence); SAM $680M (vertical-specific AI prospecting)'),
    ('Target Customer', 'VP Sales / Head of RevOps at companies selling into manufacturing: $2M–$50M ARR'),
    ('Pricing Model', '$1,500–$8,000/month SaaS; usage-based API tier for enterprises'),
    ('Valuation Range', '$8M–$25M at seed/Series A based on SaaS comps (8–15× ARR)'),
    ('Time to Market-Ready', '6–8 months to MVP launch; 12–14 months to full v1.0'),
    ('Critical Gaps', '6 critical gaps (multi-tenancy, auth, analytics, billing, CRM sync, public UX)'),
    ('Differentiator', 'Manufacturing-specific AI research depth — not generic contact enrichment'),
]
for i, (d, f) in enumerate(rows_data):
    row = table.add_row()
    row.cells[0].text = d
    row.cells[1].text = f
    if i % 2 == 0:
        set_cell_background(row.cells[0], 'F1F5F9')
        set_cell_background(row.cells[1], 'F1F5F9')
    for cell in [row.cells[0]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: MARKET OPPORTUNITY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Market Opportunity Analysis', level=1)

add_heading(doc, '2.1  The Problem ProspectIQ Solves', level=2)
add_para(doc,
    'Selling into manufacturing is uniquely hard. Unlike SaaS or finance, manufacturing '
    'companies are operationally diverse, technology-averse, and poorly covered by generic '
    'data providers. A VP of Sales at an industrial AI startup faces:')
bullets = [
    'Apollo/ZoomInfo returns thousands of contacts but no manufacturing-specific intelligence (ERP stacks, CMMS platforms, IoT maturity, maintenance approach)',
    'No trigger event detection — they don\'t know who just hired a new VP Ops or announced a plant expansion',
    'Qualification is manual — SDRs spend hours researching before writing a single email',
    'Generic sequences get ignored — manufacturing buyers spot boilerplate instantly',
    'No sector-specific scoring — a $50M food manufacturer and a $50M automotive parts maker need completely different outreach',
]
for b in bullets:
    add_bullet(doc, b)

add_para(doc,
    'ProspectIQ eliminates all five of these problems with a vertically specialized, '
    'AI-first pipeline.', space_before=6)

add_heading(doc, '2.2  Market Sizing', level=2)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(['Segment', 'Size', 'Definition', 'ProspectIQ Share']):
    hdr2[i].text = h
    for para in hdr2[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(hdr2[i], '1E293B')
    for para in hdr2[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

market_rows = [
    ('TAM', '$4.2B', 'Global B2B data, sales intelligence, intent data platforms', '—'),
    ('SAM', '$680M', 'AI-augmented prospecting tools for industrial/manufacturing verticals', '—'),
    ('SOM (3yr)', '$34M', '~5,000 companies selling into manufacturing, $6K/yr ACV', '$1–4M ARR realistic'),
]
for i, cells in enumerate(market_rows):
    row = table2.add_row()
    for j, text in enumerate(cells):
        row.cells[j].text = text
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()
add_para(doc,
    'The manufacturing sector alone comprises 250,000+ companies in the US with $100M+ '
    'revenue. These are systematically underserved by existing B2B data providers. The '
    'companies selling industrial AI, IoT, MES, CMMS, and ERP solutions into this market '
    'represent the immediate buyer base for ProspectIQ.')

add_heading(doc, '2.3  Target Customer Segments', level=2)

segments = [
    ('Primary: Industrial Tech Startups ($2M–$20M ARR)',
     'Series A/B companies selling predictive maintenance, IoT platforms, energy management, '
     'MES/ERP, or industrial AI. Have 3–15 person sales teams, need qualified pipeline fast. '
     'Pain: no vertical-specific data, generic sequences don\'t convert. WTP: $2,000–$6,000/mo.'),
    ('Secondary: Industrial Software Vendors ($20M–$200M ARR)',
     'Established players (Rockwell, Emerson resellers, CMMS vendors, ERP implementers) '
     'expanding GTM coverage. Have RevOps teams who will buy tooling if it reduces SDR '
     'research time. WTP: $6,000–$15,000/mo for team seats.'),
    ('Tertiary: Industrial Services Firms',
     'Engineering consultancies, system integrators, MRO distributors entering new verticals. '
     'Less sophisticated GTM motion but high ACV per customer. WTP: $1,500–$3,000/mo.'),
    ('API Tier: Data Consumers',
     'Companies wanting access to ProspectIQ\'s manufacturing intelligence data via API '
     '(for enriching their own CRM, powering their own models). $500–$2,000/mo by credits.'),
]
for title, desc in segments:
    add_para(doc, title, bold=True, space_before=6)
    add_para(doc, desc, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Competitive Landscape & Positioning', level=1)

add_heading(doc, '3.1  Competitive Matrix', level=2)

comp_table = doc.add_table(rows=1, cols=6)
comp_table.style = 'Table Grid'
headers = ['Platform', 'Manufacturing Depth', 'AI Research', 'Trigger Events', 'Qualification AI', 'Pricing']
for i, h in enumerate(headers):
    comp_table.rows[0].cells[i].text = h
    for para in comp_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(comp_table.rows[0].cells[i], '1E293B')
    for para in comp_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

comps = [
    ('Apollo.io', '✗ Generic', '✗ None', '✗ None', '✗ None', '$49–$149/mo'),
    ('ZoomInfo', '★ Basic NAICS', '✗ None', '★ Basic intent', '✗ None', '$15K–$40K/yr'),
    ('Cognism', '✗ Generic', '✗ None', '★ Job change alerts', '✗ None', '$1,500+/mo'),
    ('Clay.com', '✗ Generic', '★ GPT enrichment', '✗ None', '✗ None', '$149–$800/mo'),
    ('6sense', '★ NAICS filter', '✗ None', '★★ Intent signals', '✗ None', '$60K+/yr'),
    ('ProspectIQ', '★★★ Deep mfg', '★★★ Claude research', '★★★ 7 trigger types', '★★★ PQS scoring', '$1,500–$8K/mo'),
]
for i, row_data in enumerate(comps):
    row = comp_table.add_row()
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if j == 0:
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
    if i == len(comps) - 1:
        for cell in row.cells:
            set_cell_background(cell, 'DCFCE7')
    elif i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()

add_heading(doc, '3.2  ProspectIQ\'s Defensible Moat', level=2)
moat_items = [
    ('Vertical depth at the research layer',
     'Knowing that a company uses "Allen-Bradley PLCs, IBM Maximo CMMS, and recently hired '
     'a new VP of Operations from Honeywell" is not something Apollo or ZoomInfo can produce. '
     'This requires LLM-powered synthesis of multiple data sources with manufacturing-specific prompts.'),
    ('PQS Scoring Engine',
     'A structured 100-point qualification framework (firmographic + technographic + '
     'timing/pain + engagement) that replaces subjective SDR judgment with consistent, '
     'auditable scores. No competitor has this at the sector level.'),
    ('Trigger Event Detection',
     '7 buying trigger categories (leadership change, M&A/PE, ESG commitment, operational '
     'incident, CapEx, growth signal, competitor displacement) specifically tuned for '
     'manufacturing sales cycles.'),
    ('Data network effects',
     'As more users run ProspectIQ against manufacturing targets, research quality improves '
     'and the knowledge base compounds — especially if company intelligence is cached and '
     'refreshed rather than re-researched on every query.'),
]
for title, desc in moat_items:
    add_bullet(doc, f': {desc}', bold_prefix=title)

add_heading(doc, '3.3  Positioning Statement', level=2)
add_para(doc,
    '"ProspectIQ is the only AI-native sales intelligence platform built for companies '
    'selling into manufacturing. Where Apollo gives you contacts and ZoomInfo gives you '
    'company data, ProspectIQ gives you a complete picture of every manufacturing prospect: '
    'their technology stack, maintenance approach, buying triggers, and a qualified score — '
    'so your SDRs can skip the research and go straight to the conversation."',
    italic=True, space_before=4, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: PRICING & VALUATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Pricing & Valuation Model', level=1)

add_heading(doc, '4.1  Recommended Pricing Tiers', level=2)

pricing_table = doc.add_table(rows=1, cols=5)
pricing_table.style = 'Table Grid'
for i, h in enumerate(['Tier', 'Price', 'Target', 'Inclusions', 'Limits']):
    pricing_table.rows[0].cells[i].text = h
    for para in pricing_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(pricing_table.rows[0].cells[i], '1E293B')
    for para in pricing_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

tiers = [
    ('Starter', '$1,500/mo', 'Indie reps, tiny startups', '500 researched companies/mo, 1 user, basic PQS, email sequences', 'No CRM sync, no API'),
    ('Growth', '$3,500/mo', 'Series A startups, 3–10 rep teams', '2,000 companies/mo, 5 users, full PQS, CRM sync (HubSpot/Salesforce), trigger alerts', 'No white-label'),
    ('Scale', '$7,500/mo', '$20M+ ARR vendors, enterprise teams', '10,000 companies/mo, 20 users, API access, custom ICPs, dedicated CSM, Slack integration', 'Custom pricing above'),
    ('API', '$0.05/company', 'Data consumers, integrators', 'Per-enrichment credit: research + contacts + scoring', 'Min $500/mo commitment'),
]
for i, (tier, price, target, inc, lim) in enumerate(tiers):
    row = pricing_table.add_row()
    row.cells[0].text = tier
    row.cells[1].text = price
    row.cells[2].text = target
    row.cells[3].text = inc
    row.cells[4].text = lim
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()
add_para(doc,
    'Annual prepay discount: 15%. Volume pricing available above 50,000 companies/month. '
    'All plans include the Instantly.ai integration for outreach sequencing.',
    italic=True, size=9.5)

add_heading(doc, '4.2  Revenue Model & Unit Economics', level=2)

rev_items = [
    ('Cost per company researched', '$0.013 (Claude Sonnet 4.6, ~2,100 tokens avg)'),
    ('Cost per contact enriched', '$0.05–$0.10 (Apollo People Match API)'),
    ('Gross margin at Growth tier', '~78% (API costs ~$770/mo at 2K companies + enrichment)'),
    ('Blended CAC (self-serve + PLG)', '$2,500–$5,000'),
    ('Expected LTV at Growth tier', '$42,000 (12-month average retention)'),
    ('LTV:CAC ratio target', '8–12× (achievable by month 18)'),
]
rev_table = doc.add_table(rows=1, cols=2)
rev_table.style = 'Table Grid'
for i, h in enumerate(['Metric', 'Value']):
    rev_table.rows[0].cells[i].text = h
    for para in rev_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(rev_table.rows[0].cells[i], '1E293B')
    for para in rev_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
for i, (m, v) in enumerate(rev_items):
    row = rev_table.add_row()
    row.cells[0].text = m
    row.cells[1].text = v
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()

add_heading(doc, '4.3  Valuation Scenarios', level=2)

val_table = doc.add_table(rows=1, cols=4)
val_table.style = 'Table Grid'
for i, h in enumerate(['Scenario', 'ARR @ 12mo', 'Revenue Multiple', 'Valuation']):
    val_table.rows[0].cells[i].text = h
    for para in val_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(val_table.rows[0].cells[i], '1E293B')
    for para in val_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

vals = [
    ('Base (20 customers)', '$600K', '10–12×', '$6M–$7.2M'),
    ('Target (50 customers)', '$1.5M', '12–15×', '$18M–$22.5M'),
    ('Upside (120 customers + API)', '$3.8M', '15–20×', '$57M–$76M'),
]
for i, cells in enumerate(vals):
    row = val_table.add_row()
    for j, text in enumerate(cells):
        row.cells[j].text = text
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()
add_para(doc,
    'Comparable exits: Demandbase ($575M), Bombora ($100M+), G2 ($157M). '
    'Strategic acquirers include HubSpot, Salesforce, ZoomInfo, and industrial-focused '
    'PE firms. Vertical SaaS with >$1M ARR commands 12–18× revenue multiples in current market.',
    italic=True, size=9.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: GAP ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Gap Analysis: Current State vs Market-Ready', level=1)

add_para(doc,
    'ProspectIQ today is a fully functional internal tool with production-grade pipeline '
    'automation. The following analysis identifies what must be built to serve paying '
    'external customers at scale.', space_after=8)

add_heading(doc, '5.1  What Is Already Built (Strengths)', level=2)
strengths = [
    'End-to-end pipeline automation: Discovery → Research → Qualification → Enrichment → Outreach (5 stages, all functional)',
    'Deep research engine: Claude-powered single-call research extracting 15+ structured intelligence fields per company',
    '7-trigger event detection system with relevance scoring (0–10 trigger score)',
    'PQS scoring framework: 4-dimension 100-point qualification with classification buckets (qualified/high_priority/hot_prospect)',
    'Instantly.ai outreach integration with cluster × persona campaign routing (35+ sequence mappings)',
    'Supabase backend with 15+ tables, RLS, and clean data model',
    'Configurable ICP via YAML (25+ NAICS tiers, persona classification, scoring weights)',
    'Apollo.io integration for company discovery (mixed companies search) and contact enrichment (People Match)',
    'Slack notifications for pipeline events',
    'T1/T2/T3 tranche system for outreach mode control',
    'Parallel pipeline orchestration via shell scripts (research shards + qualification + enrichment loops)',
]
for s in strengths:
    add_bullet(doc, s)

add_heading(doc, '5.2  Critical Gaps (Must-Have for Market Launch)', level=2)
add_para(doc,
    'These gaps block commercial launch. None of the following exist today:', space_after=4)

critical_gaps = [
    ('GAP-1: Multi-Tenancy & User Management',
     'The system is single-user, single-workspace. There is no concept of organizations, '
     'teams, seats, or user roles. Every customer would share the same Supabase tables.',
     'Build: Supabase auth, `organizations` + `users` + `team_members` tables, '
     'row-level security per org_id, role system (admin/member/viewer). ~3 weeks.'),
    ('GAP-2: Web Application (Authenticated UI)',
     'There is no browser-accessible product. The only "UI" is a Next.js admin scaffold '
     '(limited read-only views). Customers cannot log in, view their pipeline, manage ICPs, '
     'approve outreach, or see analytics.',
     'Build: Full Next.js app with auth (Supabase Auth / Auth0), dashboard, pipeline views, '
     'company detail pages, ICP configuration, outreach approval queue. ~8 weeks.'),
    ('GAP-3: Billing & Subscription Management',
     'No payment infrastructure exists. There is no Stripe integration, no plan enforcement, '
     'no usage metering, no invoicing.',
     'Build: Stripe Billing, Stripe Customer Portal, plan tier enforcement, '
     'usage tracking (companies researched/mo), overage handling. ~2 weeks.'),
    ('GAP-4: Analytics & Reporting',
     'There are no dashboards showing pipeline health, conversion rates, cost per qualified '
     'lead, sequence performance, or ROI. Customers cannot measure value.',
     'Build: Pipeline funnel metrics, outreach performance (open/reply rates from Instantly), '
     'cost tracking (API spend per company), qualification distribution charts. ~3 weeks.'),
    ('GAP-5: CRM Bidirectional Sync',
     'ProspectIQ can push to Instantly but has no HubSpot or Salesforce integration. '
     'Qualified companies and enriched contacts should flow into the customer\'s CRM '
     'automatically, and deal stage changes should flow back.',
     'Build: HubSpot CRM API integration (companies, contacts, deals), Salesforce connector, '
     'webhook-based sync. ~4 weeks for HubSpot; +3 weeks for Salesforce.'),
    ('GAP-6: Public Landing Page & Onboarding',
     'There is no marketing website, no signup flow, no self-service onboarding, '
     'and no product tour. A customer cannot discover or try the product without '
     'direct access from the team.',
     'Build: Marketing site (Next.js), signup / onboarding wizard (ICP setup, Apollo key, '
     'Instantly key), guided first-run experience. ~3 weeks.'),
]
for title, prob, sol in critical_gaps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    add_para(doc, f'Problem: {prob}', italic=True, size=10)
    add_para(doc, f'Solution: {sol}', size=10, space_after=8)

add_heading(doc, '5.3  Important Gaps (Should-Have for Competitive Product)', level=2)

important_gaps = [
    ('GAP-7: Company Intelligence Cache & Refresh',
     'Research is re-run from scratch on every pass. There is no caching of Claude research, '
     'no staleness detection, no incremental re-research for trigger events only.'),
    ('GAP-8: LinkedIn Integration',
     'LinkedIn URLs are stored but not used for live data. No Sales Navigator API integration, '
     'no LinkedIn-based trigger event detection (new hire posts, job listings), no connection tracking.'),
    ('GAP-9: Research Quality Feedback Loop',
     'No mechanism for users to flag incorrect research, correct company data, or rate '
     'research quality. No active learning loop from user corrections.'),
    ('GAP-10: Account-Based Workflow Automation',
     'No triggered workflows: when a company hits trigger_score >= 7, no automatic escalation '
     'to Slack/email/task creation. No account owner assignment.'),
    ('GAP-11: Team Collaboration Features',
     'No internal notes, no company ownership assignment, no activity feed, '
     'no @mentions or task system within the platform.'),
    ('GAP-12: Phone Number Enrichment (Production)',
     'Apollo People Match phone enrichment exists but is gated behind async webhooks '
     'and not fully integrated into the main contact workflow.'),
]
for title, desc in important_gaps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(202, 138, 4)  # amber
    add_para(doc, desc, italic=True, size=10, space_after=6)

add_heading(doc, '5.4  Gap Priority Matrix', level=2)

gap_matrix = doc.add_table(rows=1, cols=4)
gap_matrix.style = 'Table Grid'
for i, h in enumerate(['Gap', 'Priority', 'Effort', 'Blocks']):
    gap_matrix.rows[0].cells[i].text = h
    for para in gap_matrix.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(gap_matrix.rows[0].cells[i], '1E293B')
    for para in gap_matrix.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

gap_rows = [
    ('Multi-Tenancy & Auth', 'P0 — CRITICAL', '3 weeks', 'Everything'),
    ('Web Application', 'P0 — CRITICAL', '8 weeks', 'Customer usage'),
    ('Billing', 'P0 — CRITICAL', '2 weeks', 'Revenue'),
    ('Analytics', 'P1 — HIGH', '3 weeks', 'Customer retention'),
    ('CRM Sync (HubSpot)', 'P1 — HIGH', '4 weeks', 'Enterprise sales'),
    ('Landing Page / Onboarding', 'P1 — HIGH', '3 weeks', 'Acquisition'),
    ('Intelligence Cache', 'P2 — MEDIUM', '2 weeks', 'Cost efficiency'),
    ('LinkedIn Integration', 'P2 — MEDIUM', '4 weeks', 'Data quality'),
    ('Feedback Loop', 'P2 — MEDIUM', '2 weeks', 'Quality improvement'),
    ('Workflow Automation', 'P2 — MEDIUM', '3 weeks', 'Power users'),
    ('Team Collaboration', 'P3 — NICE', '3 weeks', 'Growth tier'),
    ('Phone Enrichment', 'P3 — NICE', '1 week', 'SDR efficiency'),
]
for i, (gap, pri, eff, blocks) in enumerate(gap_rows):
    row = gap_matrix.add_row()
    row.cells[0].text = gap
    row.cells[1].text = pri
    row.cells[2].text = eff
    row.cells[3].text = blocks
    if 'P0' in pri:
        set_cell_background(row.cells[1], 'FEE2E2')
    elif 'P1' in pri:
        set_cell_background(row.cells[1], 'FEF3C7')
    elif 'P2' in pri:
        set_cell_background(row.cells[1], 'DBEAFE')
    else:
        set_cell_background(row.cells[1], 'F0FDF4')
    if i % 2 == 0:
        for j in [0, 2, 3]:
            set_cell_background(row.cells[j], 'F8FAFC')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: PRODUCT VISION & FEATURE DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Product Vision & Feature Design', level=1)

add_heading(doc, '6.1  Product Philosophy', level=2)

principles = [
    ('Intelligence-first, not data-first',
     'Every feature should surface intelligence, not just data. A company card should not '
     'show "Industry: Manufacturing" — it should show "Discrete machinery manufacturer, '
     'SAP ERP + Rockwell PLCs, reactive maintenance approach, new VP Ops from Honeywell Q1 2025."'),
    ('Pipeline as the core mental model',
     'The product lives and breathes the 5-stage pipeline. Every screen reinforces the '
     'customer\'s progress from target to conversation.'),
    ('Automation handles the tedious, humans handle the strategic',
     'Discovery, research, qualification, and first-pass enrichment are fully automated. '
     'The human touches: ICP definition, outreach approval, and relationship management.'),
    ('Honest scoring over false confidence',
     'The PQS score is transparent and auditable. Users can see exactly why a company '
     'scored 72/100. This builds trust — critical for a product claiming to replace judgment.'),
]
for title, desc in principles:
    add_bullet(doc, f': {desc}', bold_prefix=title)

add_heading(doc, '6.2  Core Feature Modules', level=2)

# ── Feature Module 1: Pipeline Dashboard ──
add_heading(doc, 'Feature Module 1: Pipeline Command Center', level=3)
add_para(doc,
    'The home screen. A single view showing the health and velocity of the entire '
    'go-to-market pipeline.', space_after=4)

feat1_items = [
    'Funnel visualization: Discovered → Researched → Qualified → Enriched → Outreach → Meeting',
    'Real-time company counts at each stage with week-over-week velocity',
    'Cost tracker: total API spend (Claude + Apollo) this month vs budget',
    'Hot Prospects feed: companies with PQS ≥ 70 or trigger_score ≥ 7 that haven\'t been sequenced',
    'Pipeline alerts: stalled research batches, enrichment failures, sequence bounce rates',
    'Daily/weekly digest email: pipeline summary, new high-priority prospects, sequence stats',
]
for item in feat1_items:
    add_bullet(doc, item)

# ── Feature Module 2: Company Intelligence ──
add_heading(doc, 'Feature Module 2: Company Intelligence Cards', level=3)
add_para(doc,
    'The core product value. Each company gets a structured intelligence profile '
    'that replaces 2–3 hours of manual SDR research.', space_after=4)

feat2_items = [
    'Header: Company name, logo, website, LinkedIn, location, estimated revenue, employee count',
    'PQS Scorecard: 100-point score with breakdown by dimension + natural-language explanation',
    'Trigger Events timeline: chronological feed of detected buying triggers with outreach relevance',
    'Manufacturing Profile: type (discrete/process/hybrid), equipment types, IoT maturity rating',
    'Technology Stack: detected ERP, CMMS, SCADA/MES, PLCs, AI/ML platforms',
    'Personalization Hooks: 3–5 ready-to-use outreach hooks specific to this company',
    'Contact roster: enriched contacts with title, email, LinkedIn, persona classification',
    'Outreach history: which sequences this company is in, open/reply status',
    'Research confidence level (high/medium/low) with last-researched timestamp',
    'Manual override: ability to correct fields, add internal notes, reassign owner',
]
for item in feat2_items:
    add_bullet(doc, item)

# ── Feature Module 3: ICP Configuration ──
add_heading(doc, 'Feature Module 3: ICP & Scoring Configuration', level=3)
add_para(doc,
    'Self-service ICP management. Replace the current YAML-based config with a UI '
    'that non-technical users can operate.', space_after=4)

feat3_items = [
    'Tier builder: drag-and-drop NAICS code selector with descriptive labels ("Industrial Machinery", "Auto Parts")',
    'Company size filters: employee count ranges, revenue bands, geography',
    'Technology inclusion/exclusion: "must have SAP or Oracle ERP", "exclude if using C3.ai"',
    'Scoring weight configuration: adjust firmographic vs technographic vs timing weights',
    'Persona configuration: which job titles map to which persona types (vp_ops, plant_manager, etc.)',
    'Sequence assignment: map (cluster × persona) → Instantly campaign ID via dropdown',
    'ICP preview: "Based on your ICP, here are 10 sample companies that would qualify"',
    'Save multiple ICP profiles for different products or campaigns',
]
for item in feat3_items:
    add_bullet(doc, item)

# ── Feature Module 4: Research Queue ──
add_heading(doc, 'Feature Module 4: Research & Qualification Queue', level=3)
add_para(doc,
    'Visibility into the automated pipeline — what\'s running, what\'s next, what failed.', space_after=4)

feat4_items = [
    'Live queue view: companies awaiting research, with tier, discovery date, expected completion',
    'Research job history: completed batches with cost, duration, success rate, average confidence',
    'Qualification queue: companies in "researched" status with current PQS distribution histogram',
    'Error log: failed research (Claude API errors, JSON parse failures) with retry controls',
    'Manual trigger: research or re-research specific companies on demand',
    'Bulk operations: select companies, bulk qualify, bulk add to sequence, bulk export',
    'Cost estimator: "Researching 500 more companies will cost ~$6.50 and take ~2 hours"',
]
for item in feat4_items:
    add_bullet(doc, item)

# ── Feature Module 5: Outreach Management ──
add_heading(doc, 'Feature Module 5: Outreach Approval & Management', level=3)
add_para(doc,
    'The human-in-the-loop layer between automated qualification and sequence launch.', space_after=4)

feat5_items = [
    'Approval queue: enriched companies ready to send to sequences, awaiting human approval',
    'Company preview: show intelligence card inline so approver has full context',
    'Approve (add to sequence), Skip (do not sequence), Watchlist, Disqualify actions',
    'Batch approval: approve all companies above PQS threshold with one click',
    'Sequence preview: which Instantly campaign this company maps to, and why',
    'Outreach tracker: active sequences per company, current step, last event (open/click/reply)',
    'Reply detection: when a prospect replies, surface it prominently with company context',
    'Meeting booked: mark companies as converted, track from which trigger/sequence',
]
for item in feat5_items:
    add_bullet(doc, item)

# ── Feature Module 6: Analytics ──
add_heading(doc, 'Feature Module 6: Analytics & Performance', level=3)
feat6_items = [
    'Funnel conversion rates: Discovered→Qualified rate, Qualified→Sequenced rate, Sequenced→Reply rate',
    'Pipeline velocity: average days from discovery to first outreach',
    'Cost per outcome: cost per qualified prospect, cost per reply, cost per meeting',
    'Research quality metrics: average confidence score, override rate (how often humans correct AI)',
    'Trigger score distribution: histogram of trigger scores across researched companies',
    'Outreach performance (from Instantly webhook): open rate, reply rate, positive reply rate by campaign',
    'ICP effectiveness: which tier + persona combinations have highest reply rates',
    'Weekly email digest with key metrics and anomaly alerts',
]
for item in feat6_items:
    add_bullet(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7: UX DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. User Experience Design', level=1)

add_heading(doc, '7.1  Design Language', level=2)
add_para(doc,
    'ProspectIQ should feel like a professional tool used by serious GTM teams — '
    'not a startup toy. The visual language mirrors Digitillis: data-dense, clean, '
    'dark-mode-first, with precise use of color to signal status and priority.')

design_principles = [
    ('Color System',
     'Primary: Teal (#0F766E) for active states, CTAs, pipeline progress. '
     'Slate (#1E293B) for dark backgrounds. '
     'Amber (#D97706) for medium-priority alerts. '
     'Red (#DC2626) for blockers/errors. '
     'Green (#16A34A) for success/hot prospects.'),
    ('Typography', 'Inter (UI), JetBrains Mono (data/scores/IDs). '
     'Hierarchy: 14px body, 12px metadata, 11px badges.'),
    ('Components', 'shadcn/ui base components, Tailwind CSS, Recharts for analytics. '
     'Consistent with Digitillis frontend stack.'),
    ('Dark Mode', 'Dark mode default (slate-900 background). '
     'Optional light mode toggle. Status badges always high-contrast.'),
    ('Data Density', 'Tables and cards are information-rich. '
     'No padding waste. Score bars, spark lines, and confidence indicators '
     'embedded inline — not hidden behind clicks.'),
]
for title, desc in design_principles:
    add_bullet(doc, f': {desc}', bold_prefix=title)

add_heading(doc, '7.2  Navigation Architecture', level=2)
add_para(doc, 'Sidebar navigation with 6 primary sections:', space_after=4)

nav_items = [
    ('Pipeline', 'The command center home — funnel overview, hot prospects, alerts'),
    ('Companies', 'Full searchable company list with filter/sort by status, PQS, tier, trigger score'),
    ('Contacts', 'Cross-company contact roster — enriched contacts, sequencing status'),
    ('Outreach', 'Approval queue, active sequences, reply tracker, meeting log'),
    ('Analytics', 'Funnel metrics, cost tracking, outreach performance'),
    ('Settings', 'ICP configuration, team management, integrations (Apollo, Instantly, HubSpot, Slack)'),
]
nav_table = doc.add_table(rows=1, cols=2)
nav_table.style = 'Table Grid'
for i, h in enumerate(['Section', 'Purpose']):
    nav_table.rows[0].cells[i].text = h
    for para in nav_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(nav_table.rows[0].cells[i], '1E293B')
    for para in nav_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
for i, (sec, purpose) in enumerate(nav_items):
    row = nav_table.add_row()
    row.cells[0].text = sec
    row.cells[1].text = purpose
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()

add_heading(doc, '7.3  Key Screen Designs', level=2)

screens = [
    ('Pipeline Command Center (Home)',
     [
         'Top row: 5 funnel stage cards (Discovered / Researched / Qualified / Enriched / Sequenced) '
         'with count, trend arrow (vs last week), and conversion rate between stages',
         'Second row: Cost tracker (Claude spend + Apollo spend this month) | Hot Prospects count | '
         'Pending approval count | Active sequences count',
         'Main section: "Hot Prospects" table — companies with PQS ≥ 70 not yet sequenced, '
         'sortable by trigger score, shows top 3 personalization hooks inline',
         'Right rail: Pipeline alerts feed (failures, stalled batches, new high-priority prospects)',
     ]),
    ('Company Detail Page',
     [
         'Sticky header: Company name + logo + website + LinkedIn icon + status badge (pill)',
         'PQS Score card: Circular score ring (0–100) with 4 dimension bars below, '
         '"Why this score" expandable explanation',
         'Trigger Events section: Timeline-style cards, each with trigger type icon, '
         'date, description, and "Outreach Relevance" text in teal italic',
         'Manufacturing Intelligence: 3-column layout — '
         'Tech Stack (chip pills) | Operations Profile (IoT maturity meter, maintenance approach) | '
         'Pain Points (bulleted, sourced from research)',
         'Personalization Hooks: Numbered list of 3–5 hooks with copy-to-clipboard button on each',
         'Contacts: Table with photo, name, title, seniority badge, email (masked unless enriched), '
         'persona badge, LinkedIn button, "Add to sequence" button',
         'Outreach History: Timeline of sequence steps, opens, clicks, replies',
         'Internal: Notes text area, owner dropdown, custom tags',
     ]),
    ('Outreach Approval Queue',
     [
         'Split-pane layout: Left = company list (sorted by PQS desc), '
         'Right = company detail inline preview',
         'Each list item shows: Company name | Tier | PQS score | Trigger score | '
         'Target sequence name | Top personalization hook',
         'Action bar: Approve (green) | Skip (grey) | Watchlist (amber) | Disqualify (red)',
         '"Approve All Above 75" bulk action button at top',
         'Approved companies immediately enter the sequence queue — status shown as "Pending sequence"',
     ]),
    ('Analytics Dashboard',
     [
         'Funnel chart: Sankey-style waterfall from Discovered to Meetings showing conversion at each stage',
         'Cost breakdown: Stacked bar chart — Claude research cost | Apollo enrichment cost — by week',
         'Outreach performance: Line chart showing open rate, reply rate, positive reply rate by week',
         'ICP Effectiveness: Heatmap of (tier × persona) → reply rate — shows which combinations convert',
         'Trigger Score Distribution: Histogram of trigger scores at time of sequencing',
         'Summary cards: This month\'s CPL (cost per qualified lead), CPM (cost per meeting), '
         'average PQS at time of outreach',
     ]),
]
for screen_title, elements in screens:
    add_para(doc, screen_title, bold=True, size=11, space_before=8, space_after=4)
    for elem in elements:
        add_bullet(doc, elem)

add_heading(doc, '7.4  Onboarding Flow', level=2)
onboarding_steps = [
    ('Step 1: Account Setup', 'Email + password or SSO. Organization name. Invite teammates.'),
    ('Step 2: ICP Definition', 'Select manufacturing sub-sectors from visual tile picker. '
     'Set company size range. Choose geographies.'),
    ('Step 3: API Keys', 'Connect Apollo.io (paste key + test). Connect Instantly.ai (paste key + test). '
     'Optionally connect HubSpot / Salesforce.'),
    ('Step 4: Persona Config', 'Select which job titles you target (VP Ops, Plant Manager, COO, etc.). '
     'Map to campaign clusters automatically.'),
    ('Step 5: First Discovery Run', '"We\'ve found 847 companies matching your ICP. '
     'Start research pipeline?" — one click to launch.'),
    ('Step 6: Sample Output', 'Show 3 fully-researched example companies from their ICP. '
     '"This is what ProspectIQ will build for every company in your pipeline."'),
]
for step, desc in onboarding_steps:
    add_bullet(doc, f': {desc}', bold_prefix=step)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8: EFFORT ESTIMATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Effort Estimation', level=1)

add_heading(doc, '8.1  Estimation Methodology', level=2)
add_para(doc,
    'Estimates use story points where 1 SP ≈ 1 engineering day (full-stack, including '
    'testing and review). Team assumption: 1 senior full-stack engineer + 1 part-time '
    'frontend/design resource. Backend leverages existing Supabase schema where possible.')

add_heading(doc, '8.2  Feature Effort Breakdown', level=2)

effort_table = doc.add_table(rows=1, cols=5)
effort_table.style = 'Table Grid'
for i, h in enumerate(['Feature', 'Phase', 'Backend (SP)', 'Frontend (SP)', 'Total (SP)']):
    effort_table.rows[0].cells[i].text = h
    for para in effort_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(effort_table.rows[0].cells[i], '1E293B')
    for para in effort_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

effort_rows = [
    # Phase 1
    ('Supabase Auth + Multi-Tenancy', '1', '10', '2', '12'),
    ('Organization & User Management', '1', '6', '5', '11'),
    ('Stripe Billing + Plan Enforcement', '1', '8', '4', '12'),
    ('Navigation Shell + Auth Pages', '1', '2', '5', '7'),
    ('Public Landing Page', '1', '1', '8', '9'),
    ('Onboarding Flow (6 steps)', '1', '4', '8', '12'),
    # Phase 2
    ('Pipeline Command Center', '2', '5', '8', '13'),
    ('Company List (filters, search, sort)', '2', '3', '6', '9'),
    ('Company Detail / Intelligence Card', '2', '2', '10', '12'),
    ('ICP Configuration UI', '2', '4', '8', '12'),
    ('Research Queue & Controls', '2', '3', '6', '9'),
    ('Outreach Approval Queue', '2', '3', '7', '10'),
    # Phase 3
    ('Analytics Dashboard', '3', '6', '10', '16'),
    ('HubSpot CRM Sync', '3', '10', '3', '13'),
    ('Reply/Meeting Tracker', '3', '4', '4', '8'),
    ('Intelligence Cache & Refresh', '3', '6', '2', '8'),
    ('Workflow Automation', '3', '8', '4', '12'),
    ('Team Collaboration (notes, owners)', '3', '4', '5', '9'),
    # Phase 4
    ('Salesforce Connector', '4', '10', '2', '12'),
    ('LinkedIn Integration (basic)', '4', '8', '3', '11'),
    ('Public API + API Keys', '4', '8', '4', '12'),
    ('Mobile-Responsive Audit', '4', '1', '6', '7'),
    ('TOTAL', '', '116', '120', '236'),
]
for i, row_data in enumerate(effort_rows):
    row = effort_table.add_row()
    for j, text in enumerate(row_data):
        row.cells[j].text = text
    if row_data[0] == 'TOTAL':
        for cell in row.cells:
            set_cell_background(cell, '1E293B')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    elif row_data[1] == '1':
        set_cell_background(row.cells[1], 'FEE2E2')
    elif row_data[1] == '2':
        set_cell_background(row.cells[1], 'FEF3C7')
    elif row_data[1] == '3':
        set_cell_background(row.cells[1], 'DBEAFE')
    elif row_data[1] == '4':
        set_cell_background(row.cells[1], 'F0FDF4')
    if i % 2 == 0 and row_data[0] != 'TOTAL':
        for j in [0, 2, 3, 4]:
            set_cell_background(row.cells[j], 'F8FAFC')

doc.add_paragraph()
add_para(doc,
    '236 story points at 1 SP/day = ~47 weeks of engineering time. '
    'With a 2-person team working in parallel (50% frontend overlap), '
    'realistic calendar time: 6 months to Phase 1+2 (MVP), '
    '12 months to full v1.0. This assumes no part-time gaps.',
    italic=True, size=9.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9: IMPLEMENTATION PLAN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Implementation Plan', level=1)

add_heading(doc, '9.1  Phased Roadmap', level=2)

phases = [
    ('Phase 1: Foundation (Weeks 1–10)',
     'P0 Gaps: Multi-tenancy, billing, public UX',
     [
         'Sprint 1–2: Supabase Auth setup, organization model, RLS policies, user management API',
         'Sprint 3–4: Stripe Billing integration, plan enforcement middleware, usage metering',
         'Sprint 5–6: Public landing page, signup flow, onboarding wizard (Steps 1–6)',
         'Sprint 7–8: Navigation shell, authenticated layout, company list page (read-only)',
         'Sprint 9–10: Beta testing with 3–5 design partners, bug fixes, performance tuning',
     ],
     'Milestone: First paying customer can log in, complete onboarding, view their pipeline',
     '~51 story points'),

    ('Phase 2: Core Product (Weeks 11–22)',
     'P1 Gaps: Web app, analytics, CRM sync',
     [
         'Sprint 11–12: Pipeline Command Center (funnel metrics, hot prospects, alerts)',
         'Sprint 13–14: Company Intelligence Card (full detail page with all fields)',
         'Sprint 15–16: ICP Configuration UI (visual tier builder, scoring weights, persona config)',
         'Sprint 17–18: Research Queue & Controls (live queue view, cost estimator, bulk ops)',
         'Sprint 19–20: Outreach Approval Queue (approval flow, sequence assignment)',
         'Sprint 21–22: Analytics Dashboard (funnel conversion, cost tracking, outreach perf)',
     ],
     'Milestone: Fully self-service product. Customer can run full pipeline end-to-end without support',
     '~65 story points'),

    ('Phase 3: Scale Features (Weeks 23–34)',
     'P2 Gaps: Intelligence quality, integrations, automation',
     [
         'Sprint 23–24: HubSpot CRM bidirectional sync (companies, contacts, deals)',
         'Sprint 25–26: Reply/meeting tracker (Instantly webhook → ProspectIQ)',
         'Sprint 27–28: Intelligence cache & refresh system (staleness detection, incremental re-research)',
         'Sprint 29–30: Workflow automation (trigger-score → Slack alert, auto-approve rules)',
         'Sprint 31–32: Team collaboration (notes, owner assignment, activity feed)',
         'Sprint 33–34: Performance hardening, mobile audit, security review',
     ],
     'Milestone: Enterprise-ready product. Viable for Series A GTM teams with 10+ seats',
     '~66 story points'),

    ('Phase 4: Platform Expansion (Weeks 35–46)',
     'P3 Gaps: API, LinkedIn, Salesforce',
     [
         'Sprint 35–36: Public API with API key management (research, qualify, enrich endpoints)',
         'Sprint 37–38: Salesforce connector (leads, contacts, accounts, opportunities)',
         'Sprint 39–40: LinkedIn integration (profile enrichment, connection tracking)',
         'Sprint 41–42: Mobile-responsive final pass, PWA support',
         'Sprint 43–44: Marketplace listing (HubSpot App Marketplace, Salesforce AppExchange)',
         'Sprint 45–46: Partner program foundations (reseller portal, white-label prep)',
     ],
     'Milestone: Full platform v1.0. API enables ecosystem integrations. Marketplace listed.',
     '~54 story points'),
]

for phase_title, subtitle, sprints, milestone, points in phases:
    p = doc.add_paragraph()
    run = p.add_run(phase_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(15, 118, 110)

    add_para(doc, subtitle, italic=True, size=10, space_after=4)
    for sprint in sprints:
        add_bullet(doc, sprint)
    add_para(doc, f'✓ {milestone}', bold=True, size=10, space_before=4)
    add_para(doc, f'Effort: {points}', size=9.5, italic=True, space_after=10)

add_heading(doc, '9.2  Technical Architecture for Market Product', level=2)

arch_items = [
    ('Backend', 'FastAPI (Python) — existing ProspectIQ patterns. '
     'Supabase PostgreSQL with RLS for multi-tenancy. '
     'Background tasks via Celery or Supabase Edge Functions.'),
    ('Frontend', 'Next.js 14 App Router + TypeScript + Tailwind CSS + shadcn/ui. '
     'Same stack as Digitillis. Dark-mode-first. Recharts for analytics.'),
    ('Auth', 'Supabase Auth (email/password + SSO via Auth0). '
     'JWT tokens, refresh token rotation, org-scoped sessions.'),
    ('Billing', 'Stripe Billing + Stripe Customer Portal. '
     'Webhook-based subscription event handling. Metered billing for API tier.'),
    ('Infrastructure', 'Railway or Vercel (frontend) + Supabase (DB + Auth). '
     'Redis for queue management and caching. Cloudflare for DNS + DDoS.'),
    ('Integrations', 'Apollo.io SDK, Instantly.ai REST API, HubSpot CRM API, '
     'Salesforce REST API, Slack webhooks, Anthropic API.'),
    ('Observability', 'PostHog (product analytics), Sentry (error tracking), '
     'Datadog or Railway metrics for infrastructure.'),
]
arch_table = doc.add_table(rows=1, cols=2)
arch_table.style = 'Table Grid'
for i, h in enumerate(['Layer', 'Technology Choice']):
    arch_table.rows[0].cells[i].text = h
    for para in arch_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(arch_table.rows[0].cells[i], '1E293B')
    for para in arch_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
for i, (layer, tech) in enumerate(arch_items):
    row = arch_table.add_row()
    row.cells[0].text = layer
    row.cells[1].text = tech
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_background(cell, 'F8FAFC')

doc.add_paragraph()

add_heading(doc, '9.3  Go-to-Market Strategy', level=2)

gtm_items = [
    ('Beta Program (Months 1–3)',
     'Recruit 5–10 design partners from the industrial tech community. '
     'Free access in exchange for weekly feedback sessions. '
     'Target: founders and VP Sales at Series A industrial AI companies.'),
    ('Paid Launch (Month 4)',
     'Growth tier ($3,500/mo) as the entry point for beta graduates. '
     'Starter tier ($1,500/mo) for self-serve from landing page. '
     'Content marketing: "How we identified 2,500 qualified manufacturing prospects for $34."'),
    ('Channels',
     'LinkedIn outreach (using ProspectIQ itself — dogfooding), '
     'Industrial tech communities (Manufacturing.news, OpsAnthology, Augury/Uptake alumni), '
     'Partnership with Instantly.ai and Apollo.io (co-marketing), '
     'G2 listing and ProductHunt launch.'),
    ('Sales Motion',
     'PLG (self-serve) for Starter tier. '
     'Sales-assisted for Growth and Scale (30-minute demo + ICP audit). '
     'Annual contracts for Scale tier with custom MSA.'),
]
for title, desc in gtm_items:
    add_bullet(doc, f': {desc}', bold_prefix=title)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10: RISK REGISTER
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Risk Register', level=1)

risk_table = doc.add_table(rows=1, cols=5)
risk_table.style = 'Table Grid'
for i, h in enumerate(['Risk', 'Likelihood', 'Impact', 'Severity', 'Mitigation']):
    risk_table.rows[0].cells[i].text = h
    for para in risk_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_background(risk_table.rows[0].cells[i], '1E293B')
    for para in risk_table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

risks = [
    ('Apollo API cost escalation', 'Medium', 'High', 'HIGH',
     'Negotiate volume pricing; build caching layer; offer plan limits'),
    ('Claude API cost per company increases', 'Low', 'High', 'MEDIUM',
     'Implement research cache; use Haiku for re-qualification of cached companies'),
    ('Apollo bans scraping / terms violation', 'Low', 'Critical', 'HIGH',
     'Use only official Apollo API endpoints; maintain ToS compliance documentation'),
    ('Research hallucination / inaccuracy', 'Medium', 'Medium', 'MEDIUM',
     'Confidence level indicator; user feedback loop; fact-verification spot checks'),
    ('ZoomInfo launches vertical AI product', 'High', 'High', 'CRITICAL',
     'Move fast; build moat in depth of research + PQS uniqueness; lock in annual customers'),
    ('Instantly.ai API deprecation', 'Low', 'Medium', 'MEDIUM',
     'Build abstraction layer; add Smartlead/Lemlist as fallback options'),
    ('Low conversion from free trial', 'Medium', 'High', 'HIGH',
     'Time-to-value < 24 hours; show researched sample companies in onboarding'),
    ('GDPR / data privacy compliance', 'Medium', 'High', 'HIGH',
     'B2B business data exempt in most jurisdictions; add data deletion API; DPA template'),
]
for i, (risk, lik, imp, sev, mit) in enumerate(risks):
    row = risk_table.add_row()
    row.cells[0].text = risk
    row.cells[1].text = lik
    row.cells[2].text = imp
    row.cells[3].text = sev
    row.cells[4].text = mit
    if sev == 'CRITICAL':
        set_cell_background(row.cells[3], 'FEE2E2')
    elif sev == 'HIGH':
        set_cell_background(row.cells[3], 'FEF3C7')
    else:
        set_cell_background(row.cells[3], 'DBEAFE')
    if i % 2 == 0:
        for j in [0, 1, 2, 4]:
            set_cell_background(row.cells[j], 'F8FAFC')

doc.add_paragraph()

# ─── CLOSING ──────────────────────────────────────────────────────────────────
add_heading(doc, 'Conclusion', level=1)
add_para(doc,
    'ProspectIQ is not a feature or a side project — it is a venture-scale product in waiting. '
    'The core intelligence engine (Claude-powered research + PQS scoring + trigger detection) '
    'represents a genuine moat that no existing B2B data provider has replicated. '
    'The six critical gaps are all buildable, well-understood engineering problems.',
    space_after=8)

add_para(doc,
    'The strategic question is not whether ProspectIQ can become a product — it clearly can. '
    'The question is timing and resource allocation. Given the competitive window '
    '(vertical AI in GTM tools is being discovered, not yet dominated), '
    'moving in the next 6 months to a design-partner beta is the right call.',
    space_after=8)

add_para(doc,
    'The recommended immediate next steps:')
next_steps = [
    'Identify 5 design partner candidates from Avanish\'s network (Series A industrial tech founders)',
    'Scope and begin Phase 1 engineering (Supabase Auth + multi-tenancy, Stripe billing)',
    'Register domain, create minimal landing page with waitlist capture',
    'Decide: build internally vs hire 1 full-stack engineer dedicated to ProspectIQ',
    'Set a 90-day milestone: first paying customer on the platform',
]
for step in next_steps:
    add_bullet(doc, step)

doc.add_paragraph()
add_para(doc,
    '─────────────────────────────────────────────────────────────────',
    size=9)
add_para(doc,
    'Document prepared by: Avanish Mehrotra | ProspectIQ Strategic Analysis',
    size=9, italic=True)
add_para(doc,
    f'Version 1.0 | {datetime.date.today().strftime("%B %Y")} | CONFIDENTIAL',
    size=9, italic=True)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = '/Users/avanish/prospectIQ/docs/ProspectIQ_Product_Design_Document.docx'
doc.save(output_path)
print(f"Document saved: {output_path}")
